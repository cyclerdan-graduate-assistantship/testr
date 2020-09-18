import os
from shutil import copyfile
from docx.shared import Pt
from docx.shared import Inches
from docx import *
import docx
from .file import File


class WordDocument(object):

    @classmethod
    def interleave_io(cls, interleave_string, prompt_string, test_case):
        result = ""
        interleave_error = False;
        input_from_keyboard = ""
        test_case = test_case.splitlines()
        test_case.reverse();
        while len(interleave_string) > 0:
            try:
                index_of_prompt_string = interleave_string.index(prompt_string)
                if len(test_case) == 0:
                    input_from_keyboard = " "
                    interleave_error = False
                else:
                    input_from_keyboard = test_case.pop()
            except ValueError as valueError:
                index_of_prompt_string = len(interleave_string) - len(prompt_string)
                input_from_keyboard = " "
                interleave_error = True
            result = result + "\n" + interleave_string[
                                     :(index_of_prompt_string + len(prompt_string))] + input_from_keyboard
            interleave_string = interleave_string[index_of_prompt_string + len(prompt_string):]
            interleave_string = interleave_string
            if interleave_error:
                # result = "ERROR -> Problem interleaving input into results, perhaps an invalid prompt string was used\n\n" + result
                if test_case:
                    result = test_case[0] + "\n\n" + result
        return result;

    @classmethod
    def assemble_feedback_document(cls, testr, directory_of_template):
        print("Working directory: " + testr.working_directory)
        created_feedback_document_successfully = True;
        try:
            copyfile(directory_of_template, testr.working_directory + '/Feedback.docx')
            feedback_document = Document(testr.working_directory + '/Feedback.docx')

            copyfile(directory_of_template, testr.working_directory + '/tmp.docx')
            tmp_document = Document(testr.working_directory + '/tmp.docx')

            style = feedback_document.styles['Normal']
            font = style.font
            font.name = 'Courier New'
            font.size = Pt(8)
            sections = feedback_document.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            for comment_file in testr.comment_files:
                feedback_document.add_paragraph(File.get_text_from_file(comment_file), None)
            feedback_document.add_page_break()

            # quiz_doc = docx.Document(testr.working_directory + '/Quiz02v1.docx')
            quiz_doc = docx.Document(testr.working_directory + '/Quiz02v2.docx')
            feedback_document.element.body.append(quiz_doc.element.body)
            feedback_document.save(testr.working_directory + '/Feedback.docx')
            tmp_document.add_page_break()

            tmp_document.add_paragraph(File.get_text_from_file(testr.path_to_source_file), None)

            tmp_document.add_page_break()
            for compile_shell_output in testr.list_of_compile_shell_outputs:
                tmp_document.add_paragraph(compile_shell_output.output, None)
            for run_shell_output in testr.list_of_run_shell_outputs:
                tmp_document.add_paragraph(run_shell_output.output, None)
            # if not testr.testr_configuration.test_input_from_cli and len(testr.testr_configuration.test_cases) > 0:
            #     os.remove(testr.working_directory + '/out.txt')
            feedback_document.element.body.append(tmp_document.element.body)
            feedback_document.save(testr.working_directory + '/Feedback.docx')
        except Exception as e:
            print(e)
            created_feedback_document_successfully = False
        return created_feedback_document_successfully

    def assemble_feedback_document_txt(testr, directory_of_template):
        created_feedback_document_successfully = True;
        try:
            feedback_document = open(testr.working_directory + '/Feedback.txt', "w")

            for comment_file in testr.comment_files:
                feedback_document.write(File.get_text_from_file(comment_file))
            feedback_document.write("\n\n")
            feedback_document.write(File.get_text_from_file(testr.path_to_source_file))
            feedback_document.write("\n\n")
            for compile_shell_output in testr.list_of_compile_shell_outputs:
                feedback_document.write(compile_shell_output.output)
            for run_shell_output in testr.list_of_run_shell_outputs:
                feedback_document.write(run_shell_output.output)
            feedback_document.close(testr.working_directory + '/Feedback.docx')
            if not testr.testr_configuration.test_input_from_cli and len(testr.testr_configuration.test_cases) > 0:
                os.remove(testr.working_directory + '/out.txt')
        except Exception:
            created_feedback_document_successfully = False
        return created_feedback_document_successfully
